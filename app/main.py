from fastapi import FastAPI, File, UploadFile, Depends, HTTPException, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
import boto3
from .database import get_db, Base, engine
from .models import Document
from .config import settings
import uuid
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime
import PyPDF2
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document as DocxDocument  # Renamed to avoid conflict with your modelimport tempfile
import os
import tempfile
from pdf2image import convert_from_path

# Add this near the top of your main.py file
DOCUMENT_TYPE_TO_SECTIONS = {
    # EB2 NIW Document Types and their mappings to letter sections
    "Resume": ["Professional_Profile", "Career_Vision", "Track_Record_of_Success"],
    "Academic_Records": ["Education_Summary", "Academic_Credentials_Summary", "Academic_Evaluation_Summary"],
    "Professional_Plan": ["Intent_to_Work_in_Field", "Career_Vision", "Proposed_Endeavor_Overview"],
    "Letters_of_Support": ["Professional_Standing", "Recognition_and_Influence", "Track_Record_of_Impact"],
    "Expert_Opinion_Letter": ["Expert_Opinion", "Expert_Reasoning_Economic_Benefit", "Expert_Endorsement"],
    "Certifications": ["Continued_Education_Certification", "Certification_Summary"],
    "Salary_History": ["Impact_of_Endeavor", "Salary_Benchmarking"],
    "Industry_Reports": ["Labor_Market_Context", "Policy_Alignment", "Field_Overview"],
    "Publications": ["Field_Contributions", "Academic_Credentials_Summary"],
    "Memberships": ["Recognition_and_Standing", "Certifications_Continued_Learning"]
}

# Define category-to-document type mappings
CATEGORY_TO_DOCUMENT_TYPES = {
    # EB2 mappings
    "01_General_Documents": ["Resume", "Professional_Plan"],
    "02_Applicant_Background": ["Resume", "Academic_Records", "Certifications"],
    "03_NIW_Criterion_1_Significant_Merit_and_Importance": ["Professional_Plan", "Industry_Reports", "Field_Overview"],
    "04_NIW_Criterion_2_Positioned_to_Advance_the_Field": ["Resume", "Professional_Plan", "Project_Examples"],
    "05_NIW_Criterion_3_Benefit_to_USA_Without_Labor_Certification": ["Expert_Opinion_Letter", "Industry_Reports"],
    "06_Letters_of_Recommendation": ["Letters_of_Support", "Expert_Opinion_Letter"],
    "07_Peer_Reviewed_Publications": ["Publications"],
    "08_Additional_Supporting_Documents": ["Memberships", "Certifications", "Salary_History"]
}

app = FastAPI()

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve index.html at root
@app.get("/")
async def read_root():
    return FileResponse("static/index.html")

class DocumentSchema(BaseModel):
    id: str
    filename: str
    s3_url: str
    case_id: str
    visa_type: str
    category: str
    uploaded_at: datetime
    extracted_text: Optional[str] = None

    class Config:
        from_attributes = True

Base.metadata.create_all(bind=engine)

s3_client = boto3.client(
    "s3",
    region_name="us-east-2",
    aws_access_key_id=settings.AWS_ACCESS_KEY,
    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
)

# Initialize Textract client
textract_client = boto3.client(
    'textract',
    region_name="us-east-2",
    aws_access_key_id=settings.AWS_ACCESS_KEY,
    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
)

def upload_to_s3(file: UploadFile, case_id: str, visa_type: str, category: str):
    try:
        print(f"Starting S3 upload process...")
        print(f"Case ID: {case_id}")
        print(f"Visa Type: {visa_type}")
        print(f"Category: {category}")
        
        file_id = str(uuid.uuid4())
        s3_key = f"raw/{case_id}/{visa_type}/{category}/{file_id}_{file.filename}"
        
        print(f"S3 Key: {s3_key}")
        
        file.file.seek(0)
        s3_client.upload_fileobj(file.file, settings.S3_BUCKET_NAME, s3_key)
        
        s3_url = f"https://{settings.S3_BUCKET_NAME}.s3.amazonaws.com/{s3_key}"
        print(f"File uploaded successfully to: {s3_url}")
        
        return s3_url
    except Exception as e:
        print(f"S3 Upload Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"S3 Upload Error: {str(e)}")



@app.post("/upload/")
async def upload_file(
    file: UploadFile = File(...),
    case_id: str = Form(...),
    visa_type: str = Form(...),
    category: str = Form(...),
    document_type: Optional[str] = Form(None),  # Added document_type as optional
    db: Session = Depends(get_db)
):
    try:
        print(f"Processing upload for case: {case_id}")

        # Read file content once
        file_bytes = await file.read()

        # Initialize variables
        num_pages = None
        extracted_text = ""
        text_extraction_status = "not_attempted"

        # Determine document type and relevant sections
        doc_type = document_type
        if not doc_type:
            # Infer document type from category
            possible_types = CATEGORY_TO_DOCUMENT_TYPES.get(category, [])
            doc_type = possible_types[0] if possible_types else "Other_Document"

        # Map document type to relevant sections
        relevant_sections = DOCUMENT_TYPE_TO_SECTIONS.get(doc_type, [])

        # Upload to S3
        file_id = str(uuid.uuid4())
        file_extension = file.filename.split(".")[-1].lower()
        s3_key = f"raw/{case_id}/{visa_type}/{category}/{file_id}_{file.filename}"

        try:
            s3_client.put_object(
                Bucket=settings.S3_BUCKET_NAME,
                Key=s3_key,
                Body=file_bytes
            )
            s3_url = f"https://{settings.S3_BUCKET_NAME}.s3.amazonaws.com/{s3_key}"
        except Exception as e:
            print(f"S3 Upload Error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"S3 Upload Error: {str(e)}")

        # Extract text using the new method
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name

            # Extract text based on file type
            if file_extension == "pdf":
                with open(temp_file_path, 'rb') as file_handle:
                    pdf_reader = PyPDF2.PdfReader(file_handle)
                    num_pages = len(pdf_reader.pages)
                    extracted_text = ''
                    for page_num in range(num_pages):
                        page = pdf_reader.pages[page_num]
                        try:
                            extracted_text += page.extract_text() + "\n"
                        except Exception as page_error:
                            print(f"Error extracting text from page {page_num}: {str(page_error)}")
                            extracted_text += f"[Error extracting page {page_num}]\n"
            elif file_extension == "docx":
                doc = DocxDocument(temp_file_path)
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            else:
                extracted_text = "Unsupported file type"

            text_extraction_status = "completed"

            # Clean up temporary file
            os.unlink(temp_file_path)

        except Exception as extract_error:
            print(f"Text Extraction Error: {str(extract_error)}")
            extracted_text = ""
            text_extraction_status = "failed"

        # Generate metadata
        document_metadata = {
            "file_size": len(file_bytes),
            "content_type": file.content_type if hasattr(file, 'content_type') else f"application/{file_extension}",
            "upload_timestamp": datetime.utcnow().isoformat(),
            "original_filename": file.filename,
            "text_extraction_status": text_extraction_status,
            "file_extension": file_extension,
            "pages": num_pages,
            "relevant_sections_mapping": {section: True for section in relevant_sections}  # Include mapping in metadata
        }

        # Create document with all fields
        document = Document(
            filename=file.filename,
            s3_url=s3_url,
            case_id=case_id,
            visa_type=visa_type,
            category=category,
            document_type=doc_type,
            relevant_sections=relevant_sections,
            extracted_text=extracted_text,
            document_metadata=document_metadata
        )

        db.add(document)
        db.commit()

        return {
            "status": "success",
            "message": f"File uploaded successfully for case {case_id}",
            "file_url": s3_url,
            "document_type": doc_type,
            "relevant_sections": relevant_sections,
            "extracted_text_status": text_extraction_status
        }
    except Exception as e:
        print(f"Upload Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents/")
async def get_documents(
    db: Session = Depends(get_db),
    case_id: Optional[str] = None,
    visa_type: Optional[str] = None
):
    query = db.query(Document)
    if case_id:
        query = query.filter(Document.case_id == case_id)
    if visa_type:
        query = query.filter(Document.visa_type == visa_type)
    return query.order_by(Document.uploaded_at.desc()).all()

@app.get("/document_types/{visa_type}")
async def get_document_types(visa_type: str):
    """Get available document types for a specific visa type"""
    if visa_type == "EB1":
        return {
            "document_types": [
                "Resume",
                "Academic_Records",
                "Professional_Plan",
                "Letters_of_Support",
                "Publications",
                "Expert_Opinion_Letter",
                "Award_Certificates",
                "Memberships"
            ]
        }
    elif visa_type == "EB2":
        return {
            "document_types": list(set(sum([types for types in CATEGORY_TO_DOCUMENT_TYPES.values()], [])))
        }
    else:
        raise HTTPException(status_code=400, detail="Invalid visa type")
        
@app.get("/categories/{visa_type}")
async def get_categories(visa_type: str):
    if visa_type == "EB1":
        return {
            "categories": [
                "A. Evidence of receipt of lesser nationally or internationally recognized prizes or awards for excellence",
                "B. Evidence of membership in associations in the field which demand outstanding achievement",
                "C. Evidence of published material about the applicant",
                "D. Evidence that the applicant has been asked to judge the work of others",
                "E. Evidence of the applicant's original scientific, scholarly contributions",
                "F. Evidence of the applicant's authorship of scholarly articles",
                "G. Evidence that the applicant's work has been displayed",
                "H. Evidence of the applicant's performance of a leading or critical role",
                "I. Evidence that the applicant commands a high salary",
                "J. Evidence of the applicant's commercial successes",
                "Letters of Support",
                "Professional Plan"
            ]
        }
    elif visa_type == "EB2":
        return {
            "categories": [
                "01_General_Documents",
                "02_Applicant_Background",
                "03_NIW_Criterion_1_Significant_Merit_and_Importance",
                "04_NIW_Criterion_2_Positioned_to_Advance_the_Field",
                "05_NIW_Criterion_3_Benefit_to_USA_Without_Labor_Certification",
                "06_Letters_of_Recommendation",
                "07_Peer_Reviewed_Publications",
                "08_Additional_Supporting_Documents"
            ]
        }
    else:
        raise HTTPException(status_code=400, detail="Invalid visa type")


@app.get("/document_types/{visa_type}")
async def get_document_types(visa_type: str):
    if visa_type == "EB1":
        return {
            "document_types": [
                "Resume",
                "Academic_Records",
                "Professional_Plan",
                "Letters_of_Support",
                "Publications",
                "Expert_Opinion_Letter",
                "Award_Certificates",
                "Memberships"
            ]
        }
    elif visa_type == "EB2":
        return {
            "document_types": [
                "Resume",
                "Academic_Records",
                "Professional_Plan",
                "Letters_of_Support",
                "Expert_Opinion_Letter",
                "Certifications",
                "Salary_History",
                "Industry_Reports",
                "Publications",
                "Memberships"
            ]
        }
    else:
        raise HTTPException(status_code=400, detail="Invalid visa type")
        
    
@app.get("/cases/{case_id}")
async def get_case_files(case_id: str, db: Session = Depends(get_db)):
    try:
        documents = db.query(Document)\
            .filter(Document.case_id == case_id)\
            .order_by(Document.uploaded_at.desc())\
            .all()
        
        if not documents:
            return {"files": []}

        files = [{
            "id": doc.id,
            "filename": doc.filename,
            "s3_url": doc.s3_url,
            "category": doc.category,
            "uploaded_at": doc.uploaded_at,
            "visa_type": doc.visa_type,
            "preview_url": f"/preview/{doc.id}",
            "download_url": f"/download/{doc.id}"
        } for doc in documents]

        return {"files": files}
    except Exception as e:
        print(f"Error fetching case files: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/files/{file_id}")
async def delete_file(file_id: str, db: Session = Depends(get_db)):
    try:
        document = db.query(Document).filter(Document.id == file_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="File not found")

        s3_key = document.s3_url.split(f"{settings.S3_BUCKET_NAME}.s3.amazonaws.com/")[1]
        try:
            s3_client.delete_object(
                Bucket=settings.S3_BUCKET_NAME,
                Key=s3_key
            )
        except Exception as e:
            print(f"S3 Delete Error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"S3 Delete Error: {str(e)}")

        db.delete(document)
        db.commit()

        return {"message": "File deleted successfully"}
    except Exception as e:
        print(f"Delete Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/documents/{case_id}")
async def delete_case(case_id: str, db: Session = Depends(get_db)):
    try:
        documents = db.query(Document).filter(Document.case_id == case_id).all()
        
        for document in documents:
            s3_key = document.s3_url.split(f"{settings.S3_BUCKET_NAME}.s3.amazonaws.com/")[1]
            try:
                s3_client.delete_object(
                    Bucket=settings.S3_BUCKET_NAME,
                    Key=s3_key
                )
            except Exception as e:
                print(f"S3 Delete Error for file {document.filename}: {str(e)}")
                
            db.delete(document)
        
        db.commit()
        return {"message": f"Case {case_id} and all associated files deleted successfully"}
    except Exception as e:
        print(f"Case Delete Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/preview/{file_id}")
async def preview_file(file_id: str, db: Session = Depends(get_db)):
    try:
        # First check if we already have extracted text in the database
        document = db.query(Document).filter(Document.id == file_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="File not found")

        # If we already have extracted text, return it
        if document.extracted_text and document.extracted_text.strip():
            return {"text": document.extracted_text}

        # Get file from S3
        s3_key = document.s3_url.split(f"{settings.S3_BUCKET_NAME}.s3.amazonaws.com/")[1]
        response = s3_client.get_object(Bucket=settings.S3_BUCKET_NAME, Key=s3_key)
        file_bytes = response['Body'].read()

        extracted_text = ""
        
        # 1. First try Textract for non-scanned PDFs and Word docs
        try:
            print("Attempting Textract extraction...")
            textract_response = textract_client.detect_document_text(
                Document={'Bytes': file_bytes}
            )
            
            for item in textract_response["Blocks"]:
                if item["BlockType"] == "LINE":
                    extracted_text += item["Text"] + "\n"
                    
        except Exception as textract_error:
            print(f"Textract failed: {str(textract_error)}")
            extracted_text = ""

        # 2. If Textract didn't work, try OCR for scanned documents
        if not extracted_text.strip():
            print("Attempting OCR extraction...")
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_input_file:
                    temp_input_path = temp_input_file.name
                    temp_input_file.write(file_bytes)

                # Convert PDF to images
                images = convert_from_path(temp_input_path)
                
                # Process each page
                for image in images:
                    # Use pytesseract to extract text from image
                    page_text = pytesseract.image_to_string(image)
                    extracted_text += page_text + "\n\n"

                # Clean up temporary file
                os.unlink(temp_input_path)

            except Exception as ocr_error:
                print(f"OCR failed: {str(ocr_error)}")
                if os.path.exists(temp_input_path):
                    os.unlink(temp_input_path)
                extracted_text = ""

        # 3. If both failed, try PyPDF2 as last resort
        if not extracted_text.strip():
            print("Attempting PyPDF2 extraction...")
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(file_bytes)
                    temp_file_path = temp_file.name
                
                with open(temp_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        try:
                            extracted_text += page.extract_text() + "\n"
                        except Exception as page_error:
                            print(f"Error on page: {str(page_error)}")
                
                os.unlink(temp_file_path)
            except Exception as pdf_error:
                print(f"PyPDF2 failed: {str(pdf_error)}")

        # Store successfully extracted text
        if extracted_text and extracted_text.strip():
            document.extracted_text = extracted_text
            db.commit()
            print("Successfully stored extracted text")
            return {"text": extracted_text}
        else:
            print("No text could be extracted using any method")
            return {"text": "No text could be extracted from this document"}

    except Exception as e:
        print(f"Preview Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{file_id}")
async def download_file(file_id: str, db: Session = Depends(get_db)):
    try:
        document = db.query(Document).filter(Document.id == file_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="File not found")

        s3_key = document.s3_url.split(f"{settings.S3_BUCKET_NAME}.s3.amazonaws.com/")[1]
        
        try:
            response = s3_client.get_object(Bucket=settings.S3_BUCKET_NAME, Key=s3_key)
            return StreamingResponse(
                response['Body'],
                media_type='application/octet-stream',
                headers={
                    'Content-Disposition': f'attachment; filename="{document.filename}"'
                }
            )
        except Exception as e:
            print(f"S3 Download Error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"S3 Download Error: {str(e)}")

    except Exception as e:
        print(f"Download Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))